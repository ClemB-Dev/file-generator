from django.shortcuts import render
from django.http import StreamingHttpResponse, HttpResponse, FileResponse
from .forms import FileFormatForm
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import locale
from reportlab.pdfgen import canvas
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph
from reportlab.lib.pagesizes import A4

locale.setlocale(locale.LC_ALL, 'fr_FR')
width, height = A4

def index(request):
    context = {
        'form': FileFormatForm(),
    }
    return render(request, './generator/index.html', context)


def build_docx_document(data):
    document = Document()
    client_data = document.add_paragraph()
    run = client_data.add_run()
    run.bold = True
    run.add_text(str(data["client_field"]))
    run.add_break()
    run = client_data.add_run()
    run.bold = False
    run.add_text('54 rue Eugène Dupuis')
    run.add_break()
    run.add_text('94000 Créteil')

    urssaf_data = document.add_paragraph()
    paragraph_format = urssaf_data.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_format.space_after = Inches(0.5)
    run = urssaf_data.add_run()
    run.bold = True
    run.add_text(str(data["urssaf_name_field"]))
    run.add_break()
    run = urssaf_data.add_run()
    run.bold = False
    run.add_text('54 rue Eugène Dupuis')
    run.add_break()
    run.add_text('94000 Créteil')

    city_and_date = document.add_paragraph()
    paragraph_format = city_and_date.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_format.space_after = Inches(1)
    run = city_and_date.add_run()
    run.add_text('{0}, le {1}'.format(
        str(data['city_field']).capitalize(),
        data['send_date_field'].strftime("%d %B %Y")))

    document.add_paragraph('Madame, Monsieur,')

    content = document.add_paragraph()
    run = content.add_run()
    run.add_text(str(data['reason_field']))

    document.add_paragraph('Restant à votre disposition pour '
                           'tout renseignement complémentaire, '
                           'nous vous prions d\'agréer, '
                           'Madame, Monsieur, nos sincères salutations.')

    return document


def build_pdf_document(data):
    main_content_style = ParagraphStyle('main content style', alignment=4)
    right_content_style = ParagraphStyle('date time content style', alignment=2)
    buffer_bytes = io.BytesIO()
    p = canvas.Canvas(buffer_bytes, pagesize=A4)
    client_paragraph = Paragraph('<b>{0}</b><br/>54 rue kreldd<br/>54 rue krdsqdsqdsqdsqdsqdqseldd<br/>Paris 75004'.format(str(data['client_field'])))
    # client_paragraph.wrapOn(p, 150, 100)
    w, h = client_paragraph.wrap(150, 100)
    client_paragraph.drawOn(p, 90, 780-h)
    # urssaf_paragraph = Paragraph('<b>{0}</b><br/>54 rue kreldd<br/>54 rue krdsqdsqdsqdsqdsqdqseldd<br/>Paris 75004'.format(str(data['urssaf_name_field'])
    # ))
    # w, h = urssaf_paragraph.wrap(150, 100)
    # urssaf_paragraph.drawOn(p, 400, 690-h)
    urssaf_paragraph = Paragraph(
        '<b>{0}</b><br/>54 rue kreldd<br/>54 rue krdsqdsqdsqdsqdsqdqseldd<br/>Paris 75004'.format(str(data['urssaf_name_field'])
        ),
        right_content_style)
    w, h = urssaf_paragraph.wrap(width-180, 100)
    urssaf_paragraph.drawOn(p, 90, 690-h)
    date_city_paragraph = Paragraph('{0}, le {1}'.format(
        str(data['city_field']).capitalize(),
        data['send_date_field'].strftime("%d %B %Y")), right_content_style)
    w, h = date_city_paragraph.wrap(width-180, 50)
    date_city_paragraph.drawOn(p, 90, 600-h)
    main_content_paragraph = Paragraph(
                           'Madame, Monsieur,<br/><br/>{0}<br/><br/>Restant à votre disposition pour '
                           'tout renseignement complémentaire, '
                           'nous vous prions d\'agréer, '
                           'Madame, Monsieur, '
                           'nos sincères salutations'.format(str(data['reason_field'])), main_content_style)
    w, h = main_content_paragraph.wrap(width-180, 600)
    main_content_paragraph.drawOn(p, 90, 550-h)
    p.save()
    buffer_bytes.seek(0)
    return buffer_bytes


def file_generator(request):
    form = FileFormatForm(request.POST)
    print(request.method == 'POST', form.is_valid())
    if request.method == 'POST' and form.is_valid():
        if 'Word' in request.POST:
            document = build_docx_document(form.cleaned_data)
            buffer_bytes = io.BytesIO()
            document.save(buffer_bytes)
            buffer_bytes.seek(0)

            response = StreamingHttpResponse(
                streaming_content=buffer_bytes,
                content_type='application/vnd.openxmlformats-officedocument'
                             '.wordprocessingml.document'
            )

            response['Content-Disposition'] = 'attachment;filename=courrier.docx'
            response["Content-Encoding"] = 'UTF-8'
            return response
        elif 'PDF' in request.POST:
            buffer_bytes = build_pdf_document(form.cleaned_data)
            response = StreamingHttpResponse(
                streaming_content=buffer_bytes,
                content_type='application/pdf'
            )

            response['Content-Disposition'] = 'attachment;filename=courrier.pdf'
            response["Content-Encoding"] = 'UTF-8'
            return response
        else:
            return HttpResponse('Something went wrong')
    else:
        return HttpResponse('Something went wrong not')
